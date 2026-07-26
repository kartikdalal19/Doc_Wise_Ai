from pathlib import Path

import fitz
from docx import Document as DocxDocument

from core.models import LoadedDocument, DocumentPage


class DocumentLoader:
    """
    Loads PDF, DOCX and TXT documents.
    """

    SUPPORTED_EXTENSIONS = [".pdf", ".docx", ".txt"]

    def load(self, file_path: str) -> LoadedDocument:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"{file_path} does not exist.")

        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {path.suffix}"
            )

        extension = path.suffix.lower()

        if extension == ".pdf":
            return self._load_pdf(path)

        elif extension == ".docx":
            return self._load_docx(path)

        elif extension == ".txt":
            return self._load_txt(path)

    # -------------------------------------------------

    def _load_pdf(self, path: Path) -> LoadedDocument:

        pdf = fitz.open(path)

        pages = []

        for page_number, page in enumerate(pdf, start=1):

            page_text = page.get_text()

            pages.append(
                DocumentPage(
                    page_number=page_number,
                    text=page_text
                )
            )

        metadata = {
            "pages": len(pdf),
            "title": pdf.metadata.get("title"),
            "author": pdf.metadata.get("author"),
        }

        pdf.close()

        return LoadedDocument(
            filename=path.name,
            file_type="PDF",
            pages=pages,
            metadata=metadata
        )

    # -------------------------------------------------

    def _load_docx(self, path: Path) -> LoadedDocument:

        doc = DocxDocument(path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        pages = [
            DocumentPage(
                page_number=1,
                text=text
            )
        ]

        metadata = {
            "paragraphs": len(doc.paragraphs)
        }

        return LoadedDocument(
            filename=path.name,
            file_type="DOCX",
            pages=pages,
            metadata=metadata
        )

    # -------------------------------------------------

    def _load_txt(self, path: Path) -> LoadedDocument:

        text = path.read_text(
            encoding="utf-8",
            errors="ignore"
        )

        pages = [
            DocumentPage(
                page_number=1,
                text=text
            )
        ]

        metadata = {
            "characters": len(text)
        }

        return LoadedDocument(
            filename=path.name,
            file_type="TXT",
            pages=pages,
            metadata=metadata
        )